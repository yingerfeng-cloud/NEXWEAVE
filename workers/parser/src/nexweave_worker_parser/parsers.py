"""Real, bounded adapters for the six M3 document types."""

from __future__ import annotations

import asyncio
import csv
import hashlib
import html
import io
import re
import zipfile
from dataclasses import dataclass
from typing import Any
from uuid import UUID

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from nexweave_contracts import (
    DocumentSegment,
    ParseFailureUnit,
    ParserCapability,
    ParseRequest,
    ParseResultManifest,
    SourceAnchor,
    canonical_parse_result_checksum,
)
from nexweave_contracts.source import ParseSecurityStats
from nexweave_contracts.source_anchor import (
    BlockLocator,
    BoundingBoxLocator,
    CharacterRangeLocator,
    PageLocator,
    TableCellLocator,
)
from nexweave_domain import BlockType, FailureScope, LocatorStatus, ParseJobStatus
from nexweave_worker_parser.protocol import (
    ParserPolicyError,
    parser_capabilities,
    probe_type,
)


@dataclass(slots=True)
class _Builder:
    request: ParseRequest
    segments: list[DocumentSegment]
    anchors: list[SourceAnchor]
    failures: list[ParseFailureUnit]
    output_chars: int = 0

    def add(
        self,
        *,
        block_type: BlockType,
        structure_path: str,
        text: str,
        locators: tuple[Any, ...],
        page: int | None = None,
        sheet: str | None = None,
        table_id: str | None = None,
        row: int | None = None,
        column: int | None = None,
    ) -> None:
        normalized = _normalize_text(text)
        if not normalized and block_type is not BlockType.PAGE_BOUNDARY:
            return
        self.output_chars += len(normalized)
        if self.output_chars > self.request.budget.max_output_chars:
            raise ParserPolicyError(
                "PARSER_RESOURCE_LIMIT_EXCEEDED", "The normalized output exceeds its budget."
            )
        if len(self.segments) >= self.request.budget.max_segments:
            raise ParserPolicyError(
                "PARSER_RESOURCE_LIMIT_EXCEEDED", "The segment count exceeds its budget."
            )
        sequence = len(self.segments)
        segment_id = _stable_uuid7(f"{self.request.parse_job_id}:segment:{sequence}")
        all_locators = (BlockLocator(block_id=str(segment_id)), *locators)
        text_checksum = _sha(normalized)
        segment = DocumentSegment(
            id=segment_id,
            source_version_id=self.request.source.source_version_id,
            parse_job_id=self.request.parse_job_id,
            sequence=sequence,
            block_type=block_type,
            structure_path=structure_path,
            normalized_text=normalized,
            text_checksum=text_checksum,
            page_number=page,
            sheet_name=sheet,
            table_id=table_id,
            row_index=row,
            column_index=column,
            locators=all_locators,
            parser_id=self.request.parser_id,
            parser_version=self.request.parser_version,
            config_checksum=self.request.config_checksum,
        )
        self.segments.append(segment)
        self.anchors.append(
            SourceAnchor(
                id=_stable_uuid7(f"{self.request.parse_job_id}:anchor:{sequence}"),
                source_version_id=self.request.source.source_version_id,
                parse_job_id=self.request.parse_job_id,
                source_checksum=self.request.source.checksum_sha256,
                excerpt_hash=text_checksum,
                locators=all_locators,
                status=LocatorStatus.VALID,
            )
        )

    def fail(
        self,
        code: str,
        scope: FailureScope,
        scope_ref: str,
        detail: str,
        *,
        retryable: bool = False,
    ) -> None:
        self.failures.append(
            ParseFailureUnit(
                id=_stable_uuid7(
                    f"{self.request.parse_job_id}:failure:{len(self.failures)}:{code}:{scope.value}:{scope_ref}"
                ),
                parse_job_id=self.request.parse_job_id,
                error_code=code,
                scope=scope,
                scope_ref=scope_ref,
                retryable=retryable,
                safe_detail=detail,
            )
        )


class DocumentParserRegistry:
    """Capability registry with no customer or domain-specific branching."""

    def capabilities(self) -> tuple[ParserCapability, ...]:
        return parser_capabilities()

    def probe(self, *, filename: str, content_type: str, content: bytes) -> ParserCapability:
        return probe_type(filename=filename, content_type=content_type, content=content)

    async def parse(self, *, request: ParseRequest, content: bytes) -> ParseResultManifest:
        return await asyncio.to_thread(self.parse_sync, request=request, content=content)

    def parse_sync(self, *, request: ParseRequest, content: bytes) -> ParseResultManifest:
        if len(content) > request.budget.max_input_bytes:
            raise ParserPolicyError(
                "PARSER_RESOURCE_LIMIT_EXCEEDED", "The Raw input exceeds the parser budget."
            )
        self.probe(
            filename=request.filename,
            content_type=request.source.content_type,
            content=content,
        )
        if _sha_bytes(content) != request.source.checksum_sha256:
            raise ParserPolicyError(
                "SOURCE_CHECKSUM_MISMATCH", "The parser input differs from the registered Raw."
            )

        builder = _Builder(request, [], [], [])
        stats = ParseSecurityStats(input_bytes=len(content))
        content_type = request.source.content_type
        if content_type == "application/pdf":
            stats = _parse_pdf(content, builder)
        elif content_type.endswith("wordprocessingml.document"):
            stats = _parse_docx(content, builder)
        elif content_type.endswith("spreadsheetml.sheet"):
            stats = _parse_xlsx(content, builder)
        elif content_type == "text/markdown":
            _parse_markdown(content, builder)
        elif content_type == "text/plain":
            _parse_text(content, builder)
        elif content_type == "text/csv":
            _parse_csv(content, builder)
        else:
            raise ParserPolicyError("SOURCE_TYPE_UNSUPPORTED", "No approved parser is available.")

        if not builder.segments and not builder.failures:
            builder.fail(
                "PARSE_RESULT_INVALID",
                FailureScope.DOCUMENT,
                "document",
                "The document contains no usable content.",
            )
        status = ParseJobStatus.SUCCEEDED
        if builder.failures and builder.segments:
            status = ParseJobStatus.PARTIAL_FAILED
        elif builder.failures:
            status = ParseJobStatus.FAILED
        result_checksum = canonical_parse_result_checksum(
            source_checksum=request.source.checksum_sha256,
            segments=builder.segments,
            anchors=builder.anchors,
            failure_units=builder.failures,
        )
        return ParseResultManifest(
            parse_job_id=request.parse_job_id,
            source_version_id=request.source.source_version_id,
            source_checksum=request.source.checksum_sha256,
            parser_id=request.parser_id,
            parser_version=request.parser_version,
            config_checksum=request.config_checksum,
            status=status,
            segments=tuple(builder.segments),
            anchors=tuple(builder.anchors),
            failure_units=tuple(builder.failures),
            security_stats=stats,
            result_checksum=result_checksum,
        )


def _parse_pdf(content: bytes, builder: _Builder) -> ParseSecurityStats:
    try:
        reader = PdfReader(io.BytesIO(content), strict=True)
    except Exception as exc:
        raise ParserPolicyError("PARSE_RESULT_INVALID", "The PDF structure is invalid.") from exc
    if reader.is_encrypted:
        raise ParserPolicyError("SOURCE_SECURITY_POLICY_FAILED", "Encrypted PDFs are not accepted.")
    if len(reader.pages) > builder.request.budget.max_pages:
        raise ParserPolicyError(
            "PARSER_RESOURCE_LIMIT_EXCEEDED", "The PDF page count exceeds its budget."
        )

    scanned_pages = 0
    for page_index, page in enumerate(reader.pages, start=1):
        builder.add(
            block_type=BlockType.PAGE_BOUNDARY,
            structure_path=f"/pages/{page_index}",
            text=f"Page {page_index}",
            locators=(PageLocator(page=page_index),),
            page=page_index,
        )
        chunks: list[tuple[str, float, float]] = []
        visitor = _pdf_text_visitor(chunks)

        try:
            page.extract_text(visitor_text=visitor)
        except Exception:
            chunks = []
        width = max(float(page.mediabox.width), 1.0)
        height = max(float(page.mediabox.height), 1.0)
        page_text = _normalize_text(" ".join(chunk[0] for chunk in chunks))
        if len(page_text) < 8:
            scanned_pages += 1
            builder.fail(
                "OCR_REQUIRED",
                FailureScope.PAGE,
                str(page_index),
                "The page has no usable text layer and requires an approved OCR provider.",
            )
            continue
        offset = 0
        for block_index, (text, x_pos, y_pos) in enumerate(chunks):
            normalized = _normalize_text(text)
            if not normalized:
                continue
            end = offset + len(normalized.encode("utf-8"))
            x = min(max(x_pos / width, 0.0), 0.99)
            y = min(max(1.0 - y_pos / height, 0.0), 0.99)
            bbox_width = min(max(len(normalized) * 0.008, 0.01), 1.0 - x)
            bbox_height = min(0.04, 1.0 - y)
            builder.add(
                block_type=BlockType.PARAGRAPH,
                structure_path=f"/pages/{page_index}/blocks/{block_index}",
                text=normalized,
                locators=(
                    PageLocator(page=page_index),
                    CharacterRangeLocator(start=offset, end=end, text_basis="normalized_utf8"),
                    BoundingBoxLocator(
                        page=page_index,
                        x=x,
                        y=y,
                        width=bbox_width,
                        height=bbox_height,
                    ),
                ),
                page=page_index,
            )
            offset = end + 1
    return ParseSecurityStats(
        input_bytes=len(content), page_count=len(reader.pages), scanned_page_count=scanned_pages
    )


def _parse_docx(content: bytes, builder: _Builder) -> ParseSecurityStats:
    zip_stats = _validate_ooxml(content, builder.request.budget.max_input_bytes)
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ParserPolicyError("PARSE_RESULT_INVALID", "The DOCX structure is invalid.") from exc
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text
        style_name = paragraph.style.name.lower() if paragraph.style is not None else ""
        block_type = BlockType.PARAGRAPH
        if style_name.startswith("heading"):
            block_type = BlockType.HEADING
        elif style_name.startswith("list"):
            block_type = BlockType.LIST
        builder.add(
            block_type=block_type,
            structure_path=f"/document/paragraphs/{index}",
            text=text,
            locators=(),
        )
    for table_index, table in enumerate(document.tables):
        table_id = f"table-{table_index}"
        builder.add(
            block_type=BlockType.TABLE,
            structure_path=f"/document/tables/{table_index}",
            text=f"Table {table_index + 1}",
            locators=(
                TableCellLocator(
                    table_id=table_id,
                    row_start=0,
                    row_end=max(len(table.rows) - 1, 0),
                    column_start=0,
                    column_end=max(len(table.columns) - 1, 0),
                ),
            ),
            table_id=table_id,
        )
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                builder.add(
                    block_type=BlockType.TABLE_CELL,
                    structure_path=f"/document/tables/{table_index}/rows/{row_index}/cells/{column_index}",
                    text=cell.text,
                    locators=(
                        TableCellLocator(
                            table_id=table_id,
                            row_start=row_index,
                            row_end=row_index,
                            column_start=column_index,
                            column_end=column_index,
                        ),
                    ),
                    table_id=table_id,
                    row=row_index,
                    column=column_index,
                )
    return ParseSecurityStats(input_bytes=len(content), **zip_stats)


def _parse_xlsx(content: bytes, builder: _Builder) -> ParseSecurityStats:
    zip_stats = _validate_ooxml(content, builder.request.budget.max_input_bytes)
    try:
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise ParserPolicyError("PARSE_RESULT_INVALID", "The XLSX structure is invalid.") from exc
    if len(workbook.sheetnames) > builder.request.budget.max_sheets:
        raise ParserPolicyError(
            "PARSER_RESOURCE_LIMIT_EXCEEDED", "The workbook sheet count exceeds its budget."
        )
    for sheet in workbook.worksheets:
        if (
            sheet.max_row > builder.request.budget.max_rows
            or sheet.max_column > builder.request.budget.max_columns
        ):
            raise ParserPolicyError(
                "PARSER_RESOURCE_LIMIT_EXCEEDED", "The worksheet dimensions exceed their budget."
            )
        table_id = f"sheet-{sheet.title}"
        builder.add(
            block_type=BlockType.TABLE,
            structure_path=f"/workbook/sheets/{_path_part(sheet.title)}",
            text=f"Sheet {sheet.title}",
            locators=(
                TableCellLocator(
                    table_id=table_id,
                    row_start=0,
                    row_end=max(sheet.max_row - 1, 0),
                    column_start=0,
                    column_end=max(sheet.max_column - 1, 0),
                ),
            ),
            sheet=sheet.title,
            table_id=table_id,
        )
        for row_index, row in enumerate(sheet.iter_rows(values_only=True)):
            for column_index, value in enumerate(row):
                if value is None:
                    continue
                builder.add(
                    block_type=BlockType.TABLE_CELL,
                    structure_path=f"/workbook/sheets/{_path_part(sheet.title)}/rows/{row_index}/cells/{column_index}",
                    text=str(value),
                    locators=(
                        TableCellLocator(
                            table_id=table_id,
                            row_start=row_index,
                            row_end=row_index,
                            column_start=column_index,
                            column_end=column_index,
                        ),
                    ),
                    sheet=sheet.title,
                    table_id=table_id,
                    row=row_index,
                    column=column_index,
                )
    workbook.close()
    return ParseSecurityStats(
        input_bytes=len(content), sheet_count=len(workbook.sheetnames), **zip_stats
    )


def _parse_markdown(content: bytes, builder: _Builder) -> None:
    text = _decode_text(content)
    paragraph: list[tuple[str, int, int]] = []

    def flush() -> None:
        if paragraph:
            value = " ".join(item[0] for item in paragraph)
            start, end = paragraph[0][1], paragraph[-1][2]
            builder.add(
                block_type=BlockType.PARAGRAPH,
                structure_path=f"/markdown/blocks/{len(builder.segments)}",
                text=value,
                locators=(_source_range(text, start, end),),
            )
            paragraph.clear()

    cursor = 0
    for raw_line in text.splitlines(keepends=True):
        line = raw_line.rstrip("\r\n")
        line_start = cursor
        cursor += len(raw_line)
        if match := re.match(r"^(#{1,6})\s+(.+)$", line):
            flush()
            builder.add(
                block_type=BlockType.HEADING,
                structure_path=f"/markdown/headings/{len(builder.segments)}",
                text=match.group(2),
                locators=(
                    _source_range(
                        text,
                        line_start + match.start(2),
                        line_start + match.end(2),
                    ),
                ),
            )
        elif match := re.match(r"^\s*(?:[-*+] |\d+\. )(.+)$", line):
            flush()
            value = match.group(1)
            builder.add(
                block_type=BlockType.LIST,
                structure_path=f"/markdown/lists/{len(builder.segments)}",
                text=value,
                locators=(
                    _source_range(
                        text,
                        line_start + match.start(1),
                        line_start + match.end(1),
                    ),
                ),
            )
        elif line.strip():
            leading = len(line) - len(line.lstrip())
            trailing = len(line.rstrip())
            paragraph.append((line.strip(), line_start + leading, line_start + trailing))
        else:
            flush()
    flush()


def _parse_text(content: bytes, builder: _Builder) -> None:
    text = _decode_text(content)
    paragraphs = re.finditer(r"\S(?:.*?\S)?(?=\n\s*\n|\Z)", text, flags=re.DOTALL)
    for index, match in enumerate(paragraphs):
        value = _normalize_text(match.group(0))
        if value:
            builder.add(
                block_type=BlockType.PARAGRAPH,
                structure_path=f"/text/paragraphs/{index}",
                text=value,
                locators=(_source_range(text, match.start(), match.end()),),
            )


def _source_range(text: str, start: int, end: int) -> CharacterRangeLocator:
    return CharacterRangeLocator(
        start=len(text[:start].encode("utf-8")),
        end=len(text[:end].encode("utf-8")),
        text_basis="source_utf8",
    )


def _parse_csv(content: bytes, builder: _Builder) -> None:
    rows = list(csv.reader(io.StringIO(_decode_text(content), newline="")))
    if len(rows) > builder.request.budget.max_rows:
        raise ParserPolicyError(
            "PARSER_RESOURCE_LIMIT_EXCEEDED", "The CSV row count exceeds its budget."
        )
    max_columns = max((len(row) for row in rows), default=0)
    if max_columns > builder.request.budget.max_columns:
        raise ParserPolicyError(
            "PARSER_RESOURCE_LIMIT_EXCEEDED", "The CSV column count exceeds its budget."
        )
    table_id = "csv-table-0"
    builder.add(
        block_type=BlockType.TABLE,
        structure_path="/csv/table",
        text="CSV table",
        locators=(
            TableCellLocator(
                table_id=table_id,
                row_start=0,
                row_end=max(len(rows) - 1, 0),
                column_start=0,
                column_end=max(max_columns - 1, 0),
            ),
        ),
        table_id=table_id,
    )
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            builder.add(
                block_type=BlockType.TABLE_CELL,
                structure_path=f"/csv/rows/{row_index}/cells/{column_index}",
                text=value,
                locators=(
                    TableCellLocator(
                        table_id=table_id,
                        row_start=row_index,
                        row_end=row_index,
                        column_start=column_index,
                        column_end=column_index,
                    ),
                ),
                table_id=table_id,
                row=row_index,
                column=column_index,
            )


def _validate_ooxml(content: bytes, max_input_bytes: int) -> dict[str, int]:
    try:
        archive = zipfile.ZipFile(io.BytesIO(content))
        entries = archive.infolist()
    except (OSError, zipfile.BadZipFile) as exc:
        raise ParserPolicyError("PARSE_RESULT_INVALID", "The OOXML container is invalid.") from exc
    if len(entries) > 10_000:
        raise ParserPolicyError(
            "SOURCE_SECURITY_POLICY_FAILED", "The OOXML entry limit was exceeded."
        )
    normalized_names = [entry.filename.replace("\\", "/") for entry in entries]
    if len(set(normalized_names)) != len(normalized_names):
        raise ParserPolicyError(
            "SOURCE_SECURITY_POLICY_FAILED", "Duplicate OOXML archive paths are forbidden."
        )
    if any(
        name.startswith("/") or any(part == ".." for part in name.split("/"))
        for name in normalized_names
    ):
        raise ParserPolicyError(
            "SOURCE_SECURITY_POLICY_FAILED", "OOXML archive path traversal is forbidden."
        )
    if any(entry.flag_bits & 0x1 for entry in entries):
        raise ParserPolicyError(
            "SOURCE_SECURITY_POLICY_FAILED", "Encrypted OOXML entries are forbidden."
        )
    if any(
        entry.compress_type not in {zipfile.ZIP_STORED, zipfile.ZIP_DEFLATED} for entry in entries
    ):
        raise ParserPolicyError(
            "SOURCE_SECURITY_POLICY_FAILED", "Unsupported OOXML compression is forbidden."
        )
    if any(entry.file_size > max(max_input_bytes * 4, 67_108_864) for entry in entries):
        raise ParserPolicyError(
            "SOURCE_SECURITY_POLICY_FAILED", "An OOXML entry exceeded its expansion budget."
        )
    expanded_bytes = sum(entry.file_size for entry in entries)
    compressed_bytes = max(sum(entry.compress_size for entry in entries), 1)
    if (
        expanded_bytes > max(max_input_bytes * 8, 268_435_456)
        or expanded_bytes / compressed_bytes > 100
    ):
        raise ParserPolicyError(
            "SOURCE_SECURITY_POLICY_FAILED", "The OOXML expansion policy was exceeded."
        )
    names = {name.lower() for name in normalized_names}
    if any(
        "vbaproject" in name or name.endswith((".bin", ".exe", ".js", ".zip", ".7z", ".rar"))
        for name in names
    ):
        raise ParserPolicyError(
            "SOURCE_SECURITY_POLICY_FAILED", "OOXML active or embedded content is forbidden."
        )
    external = 0
    for entry in entries:
        if entry.filename.lower().endswith(".rels"):
            raw = archive.read(entry)
            external += raw.count(b'TargetMode="External"') + raw.count(b"TargetMode='External'")
    if external:
        raise ParserPolicyError(
            "SOURCE_SECURITY_POLICY_FAILED", "OOXML external relationships are forbidden."
        )
    archive.close()
    return {
        "zip_entry_count": len(entries),
        "expanded_bytes": expanded_bytes,
        "external_relationships_blocked": external,
    }


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ParserPolicyError("PARSE_RESULT_INVALID", "Text inputs must be valid UTF-8.") from exc


def _stable_uuid7(seed: str) -> UUID:
    """Derive a UUIDv7-shaped identifier so Activity retries produce the same manifest."""

    value = bytearray(hashlib.sha256(seed.encode("utf-8")).digest()[:16])
    value[6] = (value[6] & 0x0F) | 0x70
    value[8] = (value[8] & 0x3F) | 0x80
    return UUID(bytes=bytes(value))


def _pdf_text_visitor(
    target: list[tuple[str, float, float]],
) -> Any:
    def visitor(text: str, _cm: Any, tm: list[float], _font: Any, _size: float) -> None:
        if text.strip():
            target.append((text, float(tm[4]), float(tm[5])))

    return visitor


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", html.unescape(value)).strip()


def _sha(value: str) -> str:
    return f"sha256:{hashlib.sha256(value.encode()).hexdigest()}"


def _sha_bytes(value: bytes) -> str:
    return f"sha256:{hashlib.sha256(value).hexdigest()}"


def _path_part(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]", "_", value)[:255]
