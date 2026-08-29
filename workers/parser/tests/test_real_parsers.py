import hashlib
import io
import zipfile

import pytest
from docx import Document
from openpyxl import Workbook
from pydantic import ValidationError
from pypdf import PdfWriter

from nexweave_contracts import ControlledObjectRef, ParseRequest, ParseResultManifest
from nexweave_contracts.source_anchor import CharacterRangeLocator
from nexweave_domain import BlockType, ParseJobStatus, new_uuid7
from nexweave_worker_parser import DocumentParserRegistry
from nexweave_worker_parser.parsers import _validate_ooxml
from nexweave_worker_parser.protocol import ParserPolicyError


def _request(filename: str, content_type: str, content: bytes) -> ParseRequest:
    checksum = f"sha256:{hashlib.sha256(content).hexdigest()}"
    tenant, space, source, version, job = (new_uuid7() for _ in range(5))
    return ParseRequest(
        parse_job_id=job,
        source=ControlledObjectRef(
            source_version_id=version,
            object_key=f"raw/v1/{tenant}/{space}/{source}/{version}/{checksum.removeprefix('sha256:')}",
            object_version_id="v1",
            checksum_sha256=checksum,
            content_type=content_type,
            size=len(content),
        ),
        filename=filename,
        parser_id="nexweave.parser.builtin",
        parser_version="1.0.0",
        config_checksum="sha256:" + "d" * 64,
    )


@pytest.mark.parametrize(
    ("filename", "content_type", "content", "block_type"),
    [
        ("readme.md", "text/markdown", b"# Heading\n\nA paragraph.", BlockType.HEADING),
        ("notes.txt", "text/plain", b"First paragraph.\n\nSecond.", BlockType.PARAGRAPH),
        ("data.csv", "text/csv", b"name,value\nalpha,1\n", BlockType.TABLE_CELL),
    ],
)
async def test_real_text_parsers(
    filename: str, content_type: str, content: bytes, block_type: BlockType
) -> None:
    result = await DocumentParserRegistry().parse(
        request=_request(filename, content_type, content), content=content
    )
    assert result.status is ParseJobStatus.SUCCEEDED
    assert any(segment.block_type is block_type for segment in result.segments)
    assert all(anchor.parse_job_id == result.parse_job_id for anchor in result.anchors)


async def test_real_docx_parser_preserves_paragraphs_and_tables() -> None:
    document = Document()
    document.add_heading("Synthetic handbook", level=1)
    document.add_paragraph("Safety first.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "item"
    table.cell(0, 1).text = "value"
    output = io.BytesIO()
    document.save(output)
    content = output.getvalue()
    result = await DocumentParserRegistry().parse(
        request=_request(
            "handbook.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            content,
        ),
        content=content,
    )
    assert result.status is ParseJobStatus.SUCCEEDED
    assert {BlockType.HEADING, BlockType.TABLE_CELL}.issubset(
        {segment.block_type for segment in result.segments}
    )


async def test_real_xlsx_parser_preserves_sheet_and_cell_coordinates() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Synthetic"
    sheet.append(["name", "value"])
    sheet.append(["alpha", 1])
    output = io.BytesIO()
    workbook.save(output)
    content = output.getvalue()
    result = await DocumentParserRegistry().parse(
        request=_request(
            "data.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            content,
        ),
        content=content,
    )
    cells = [segment for segment in result.segments if segment.block_type is BlockType.TABLE_CELL]
    assert result.status is ParseJobStatus.SUCCEEDED
    assert cells and cells[0].sheet_name == "Synthetic"
    assert cells[0].row_index == 0 and cells[0].column_index == 0


async def test_scanned_pdf_detection_is_real_and_reports_ocr_required() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    output = io.BytesIO()
    writer.write(output)
    content = output.getvalue()
    result = await DocumentParserRegistry().parse(
        request=_request("scan.pdf", "application/pdf", content), content=content
    )
    assert result.status is ParseJobStatus.PARTIAL_FAILED
    assert result.security_stats.scanned_page_count == 1
    assert [unit.error_code for unit in result.failure_units] == ["OCR_REQUIRED"]
    assert any(segment.block_type is BlockType.PAGE_BOUNDARY for segment in result.segments)


async def test_parser_manifest_is_deterministic_for_activity_retry() -> None:
    content = b"Stable paragraph.\n\nSecond paragraph."
    request = _request("stable.txt", "text/plain", content)
    parser = DocumentParserRegistry()

    first = await parser.parse(request=request, content=content)
    replayed = await parser.parse(request=request, content=content)

    assert replayed.model_dump(mode="json") == first.model_dump(mode="json")
    assert replayed.result_checksum == first.result_checksum


async def test_text_locator_uses_global_utf8_byte_offsets() -> None:
    content = "甲段。\n\nSecond paragraph.".encode()
    result = await DocumentParserRegistry().parse(
        request=_request("offsets.txt", "text/plain", content), content=content
    )
    second = result.segments[1]
    locator = next(item for item in second.locators if isinstance(item, CharacterRangeLocator))

    assert locator.start == len("甲段。\n\n".encode())
    assert locator.end == len(content)


async def test_empty_document_fails_without_persistable_output() -> None:
    content = b"   \n\n"
    result = await DocumentParserRegistry().parse(
        request=_request("empty.txt", "text/plain", content), content=content
    )

    assert result.status is ParseJobStatus.FAILED
    assert not result.segments and not result.anchors
    assert [unit.error_code for unit in result.failure_units] == ["PARSE_RESULT_INVALID"]


async def test_manifest_rejects_tampered_result_checksum() -> None:
    content = b"Trusted manifest."
    result = await DocumentParserRegistry().parse(
        request=_request("manifest.txt", "text/plain", content), content=content
    )

    with pytest.raises(ValidationError, match="canonical manifest"):
        ParseResultManifest.model_validate(
            {**result.model_dump(mode="json"), "result_checksum": "sha256:" + "0" * 64}
        )


def test_ooxml_policy_rejects_archive_path_traversal() -> None:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr("../outside.xml", "not allowed")

    with pytest.raises(ParserPolicyError, match="path traversal"):
        _validate_ooxml(output.getvalue(), 1_024)
