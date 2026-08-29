"""Exercise the real M3 API→RustFS→ClamAV→Temporal→Parser→PostgreSQL chain."""

from __future__ import annotations

import argparse
import hashlib
import io
import subprocess
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4

import httpx
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

ROOT = Path(__file__).resolve().parents[1]
TERMINAL_PARSE = {"SUCCEEDED", "PARTIAL_FAILED", "FAILED", "CANCELED"}


def _sha(content: bytes) -> str:
    return f"sha256:{hashlib.sha256(content).hexdigest()}"


def _synthetic_docx() -> bytes:
    document = Document()
    document.add_heading("Synthetic M3 handbook", level=1)
    document.add_paragraph("This fixture contains no customer or production data.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "item"
    table.cell(0, 1).text = "value"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _synthetic_xlsx() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Synthetic"
    sheet.append(["item", "value"])
    sheet.append(["alpha", 1])
    output = io.BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _synthetic_scanned_pdf() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


@dataclass
class Api:
    client: httpx.Client
    token: str

    def request(
        self,
        method: str,
        path: str,
        *,
        expected: int = 200,
        idempotency_key: str | None = None,
        version: int | None = None,
        extra_headers: dict[str, str] | None = None,
        **kwargs: Any,
    ) -> httpx.Response:
        trace_id, span_id = uuid4().hex, uuid4().hex[:16]
        headers = {
            "Authorization": f"Bearer {self.token}",
            "traceparent": f"00-{trace_id}-{span_id}-01",
        }
        if idempotency_key:
            headers["Idempotency-Key"] = idempotency_key
        if version is not None:
            headers["If-Match"] = f'"v{version}"'
        headers.update(extra_headers or {})
        response = self.client.request(method, path, headers=headers, **kwargs)
        if response.status_code != expected:
            raise RuntimeError(
                f"{method} {path} returned {response.status_code}, expected {expected}: "
                f"{response.text[:400]}"
            )
        if response.headers.get("X-Trace-Id") != trace_id:
            raise RuntimeError(f"{method} {path} did not preserve W3C trace context")
        return response

    def wait_parse(self, parse_job_id: str, *, timeout: float = 90) -> dict[str, Any]:
        deadline = time.monotonic() + timeout
        latest: dict[str, Any] = {}
        while time.monotonic() < deadline:
            latest = self.request("GET", f"/parse-jobs/{parse_job_id}").json()
            if latest["status"] in TERMINAL_PARSE:
                return latest
            time.sleep(0.25)
        raise RuntimeError(f"ParseJob did not finish: {latest}")

    def upload(
        self,
        space_id: str,
        *,
        filename: str,
        content_type: str,
        content: bytes,
        batch_id: str,
        suffix: str,
    ) -> tuple[dict[str, Any], dict[str, Any]]:
        checksum = _sha(content)
        session = self.request(
            "POST",
            f"/spaces/{space_id}/sources/uploads",
            expected=201,
            idempotency_key=f"upload-{suffix}-{filename}",
            json={
                "filename": filename,
                "content_type": content_type,
                "expected_size": len(content),
                "expected_checksum": checksum,
                "display_name": f"Synthetic M3 {filename}",
                "description": "Synthetic M3 verification fixture",
                "classification": "INTERNAL",
                "tags": ["synthetic", "m3-e2e"],
                "import_batch_id": batch_id,
            },
        ).json()
        self.request(
            "PUT",
            f"/sources/uploads/{session['id']}/content",
            extra_headers={"Content-Type": content_type},
            content=content,
        )
        complete_key = f"complete-{suffix}-{filename}"
        completed = self.request(
            "POST",
            f"/sources/uploads/{session['id']}/complete",
            expected=202,
            idempotency_key=complete_key,
            json={"checksum": checksum, "size": len(content)},
        )
        replayed = self.request(
            "POST",
            f"/sources/uploads/{session['id']}/complete",
            expected=202,
            idempotency_key=complete_key,
            json={"checksum": checksum, "size": len(content)},
        )
        if completed.json()["parse_job_id"] != replayed.json()["parse_job_id"]:
            raise RuntimeError("idempotent complete created a second ParseJob")
        return completed.json(), self.wait_parse(completed.json()["parse_job_id"])


def _database_fact_count(event_family: str) -> int:
    env: dict[str, str] = {}
    for line in (ROOT / ".env").read_text(encoding="utf-8").splitlines():
        if line and not line.startswith("#") and "=" in line:
            key, value = line.split("=", 1)
            env[key] = value
    queries = {
        "source": (
            "SELECT count(*) FROM outbox_events WHERE event_type LIKE 'io.nexweave.source.%';"
        ),
        "parse": (
            "SELECT count(*) FROM outbox_events WHERE event_type LIKE 'io.nexweave.parse.%';"
        ),
    }
    result = subprocess.run(  # noqa: S603 - fixed local Compose verification command
        [  # noqa: S607 - reviewed fixed Docker executable
            "docker",
            "compose",
            "exec",
            "-T",
            "postgres",
            "psql",
            "-U",
            env["NEXWEAVE_POSTGRES_USER"],
            "-d",
            env["NEXWEAVE_POSTGRES_DB"],
            "-At",
            "-v",
            "ON_ERROR_STOP=1",
            "-c",
            queries[event_family],
        ],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
    )
    return int(result.stdout.strip())


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--base-url", default="http://127.0.0.1:8080/api/v1")
    args = parser.parse_args()
    suffix = uuid4().hex[:12]
    with httpx.Client(base_url=args.base_url, timeout=120) as client:
        session = client.post("/auth/dev/session", json={"subject": "local-admin"})
        session.raise_for_status()
        api = Api(client, str(session.json()["access_token"]))
        organization = api.request("GET", "/organizations").json()["items"][0]
        space = api.request(
            "POST",
            "/spaces",
            expected=201,
            idempotency_key=f"m3-space-{suffix}",
            json={
                "organization_id": organization["id"],
                "slug": f"m3-e2e-{suffix}",
                "display_name": f"M3 E2E {suffix}",
                "description": "Synthetic M3 verification space",
                "default_classification": "INTERNAL",
            },
        ).json()
        batch = api.request(
            "POST",
            f"/spaces/{space['id']}/source-import-batches",
            expected=201,
            idempotency_key=f"m3-batch-{suffix}",
            json={"display_name": f"M3 batch {suffix}"},
        ).json()
        fixtures = [
            ("notes.txt", "text/plain", b"Synthetic M3 text.\n\nSecond paragraph."),
            (
                "handbook.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                _synthetic_docx(),
            ),
            (
                "table.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                _synthetic_xlsx(),
            ),
            ("scan.pdf", "application/pdf", _synthetic_scanned_pdf()),
        ]
        results: dict[str, tuple[dict[str, Any], dict[str, Any]]] = {}
        for filename, content_type, content in fixtures:
            results[filename] = api.upload(
                str(space["id"]),
                filename=filename,
                content_type=content_type,
                content=content,
                batch_id=str(batch["id"]),
                suffix=suffix,
            )
        if results["notes.txt"][1]["status"] != "SUCCEEDED":
            raise RuntimeError("TXT parse did not succeed")
        if results["handbook.docx"][1]["status"] != "SUCCEEDED":
            raise RuntimeError("DOCX parse did not succeed")
        if results["table.xlsx"][1]["status"] != "SUCCEEDED":
            raise RuntimeError("XLSX parse did not succeed")
        scanned = results["scan.pdf"][1]
        if scanned["status"] != "PARTIAL_FAILED" or not any(
            item["error_code"] == "OCR_REQUIRED" for item in scanned["failure_units"]
        ):
            raise RuntimeError("scanned PDF did not honestly report OCR_REQUIRED")

        txt_complete = results["notes.txt"][0]
        version_id = txt_complete["source_version_id"]
        preview = api.request("GET", f"/source-versions/{version_id}/preview").json()
        if "Synthetic M3 text" not in preview["sanitized_content"]:
            raise RuntimeError("authorized text preview did not contain parsed content")
        version_response = api.request(
            "GET",
            f"/sources/{txt_complete['source_id']}/versions/{version_id}",
        )
        reparse = api.request(
            "POST",
            f"/source-versions/{version_id}/parse",
            expected=202,
            idempotency_key=f"m3-reparse-{suffix}",
            version=int(version_response.json()["version"]),
            json={
                "parser_id": "nexweave.parser.builtin",
                "parser_version": "1.0.0",
                "config": {},
            },
        ).json()
        if api.wait_parse(reparse["id"])["status"] != "SUCCEEDED":
            raise RuntimeError("reparse did not succeed")
        refreshed = api.request(
            "GET", f"/sources/{txt_complete['source_id']}/versions/{version_id}"
        ).json()
        api.request(
            "POST",
            f"/source-versions/{version_id}/invalidate",
            expected=201,
            idempotency_key=f"m3-invalidate-{suffix}",
            version=int(refreshed["version"]),
            json={
                "reason_code": "E2E_SYNTHETIC_WITHDRAWAL",
                "reason": "Synthetic M3 invalidation verification",
                "policy_version": "m3-v1",
            },
        )
        api.request("GET", f"/source-versions/{version_id}/segments", expected=409)
        masked = api.request("GET", f"/source-versions/{version_id}/preview").json()
        if masked["sanitized_content"]:
            raise RuntimeError("invalidated preview was not masked")
        batch_result = api.request("GET", f"/source-import-batches/{batch['id']}").json()
        if len(batch_result["items"]) != len(fixtures):
            raise RuntimeError("batch did not expose each file result")

    if _database_fact_count("source") < 5:
        raise RuntimeError("Source transactional Outbox facts are missing")
    if _database_fact_count("parse") < 5:
        raise RuntimeError("Parse transactional Outbox facts are missing")
    print(
        "M3 real chain verified: TXT/DOCX/XLSX, scanned PDF OCR_REQUIRED, "
        "idempotent complete, reparse, invalidation, batch and transactional Outbox."
    )
    print(
        "Not covered by this script: real OCR (not configured), broker publication, "
        "production HA/DR, parser cancellation/crash-window recovery, replay of an archived "
        "accepted M2 history, multi-architecture image promotion and external CI."
    )


if __name__ == "__main__":
    main()
